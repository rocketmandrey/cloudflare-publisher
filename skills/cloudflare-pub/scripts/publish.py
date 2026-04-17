#!/usr/bin/env python3
"""
Cloudflare Pages Publisher — публикация любого контента на Cloudflare Pages.

Вход:  .docx | .txt | .md | .html | stdin (текст/html)
Выход: публичная ссылка https://<name>.pages.dev

Требования:
    - wrangler CLI (npm i -g wrangler)
    - python-docx (pip install python-docx) — только для .docx
    - CF_ACCOUNT_ID и CF_API_TOKEN в .env

Использование:
    python3 publish.py "отчёт.docx"
    python3 publish.py "отчёт.docx" --name my-report --title "Мой отчёт"
    python3 publish.py "страница.html"
    python3 publish.py "заметки.md"
    echo "Hello world" | python3 publish.py --name test-page --stdin
    python3 publish.py "файл.txt" --html-only
"""

import sys
import os
import re
import json
import shutil
import subprocess
import tempfile
import argparse
from pathlib import Path
from datetime import datetime


SCRIPT_DIR = Path(__file__).resolve().parent
PRETTY_CSS = SCRIPT_DIR / "pretty.css"
PRETTY_TEMPLATE = SCRIPT_DIR / "pretty_template.html"


# ─── Config ───

def load_env():
    """Load .env from multiple locations (first found wins per key)."""
    env = {}
    for env_path in [
        Path(__file__).resolve().parent / ".env",
        Path.home() / ".claude" / "cloudflare-pub" / ".env",
        Path.home() / "Documents" / "personal_ai" / ".env",
        Path.cwd() / ".env",
    ]:
        if env_path.exists():
            for line in env_path.read_text().splitlines():
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                if "=" in line:
                    k, v = line.split("=", 1)
                    key = k.strip()
                    if key not in env:
                        env[key] = v.strip().strip("'\"")
    return env


ENV = load_env()
CF_ACCOUNT_ID = ENV.get("CF_ACCOUNT_ID", "")
CF_API_TOKEN = ENV.get("CF_API_TOKEN", "")


# ─── Parsers ───

def parse_docx(filepath):
    """Parse .docx → structured blocks [(type, content), ...]."""
    from docx import Document
    doc = Document(filepath)
    blocks = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            for p in doc.paragraphs:
                if p._element is element:
                    style = p.style.name if p.style else ""
                    text = p.text.strip()
                    if not text:
                        break
                    if "Heading 1" in style:
                        blocks.append(("h1", text))
                    elif "Heading 2" in style:
                        blocks.append(("h2", text))
                    elif "Heading 3" in style:
                        blocks.append(("h3", text))
                    else:
                        blocks.append(("p", text))
                    break

        elif tag == "tbl":
            for t in doc.tables:
                if t._element is element:
                    rows = []
                    for row in t.rows:
                        rows.append([cell.text.strip() for cell in row.cells])
                    blocks.append(("table", rows))
                    break

    return blocks


def parse_text(text):
    """Parse plain text / markdown → structured blocks."""
    lines = text.split("\n")
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Tab-separated table (3+ columns)
        if line.count("\t") >= 2:
            table_rows = []
            while i < len(lines) and lines[i].count("\t") >= 2:
                table_rows.append(lines[i].split("\t"))
                i += 1
            blocks.append(("table", table_rows))
            continue

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Markdown headers
        if stripped.startswith("### "):
            blocks.append(("h3", stripped[4:]))
        elif stripped.startswith("## "):
            blocks.append(("h2", stripped[3:]))
        elif stripped.startswith("# "):
            blocks.append(("h1", stripped[2:]))
        elif len(stripped) < 80 and not stripped.endswith(".") and not blocks:
            blocks.append(("h1", stripped))
        else:
            blocks.append(("p", stripped))

        i += 1

    return blocks


def parse_file(filepath):
    """Auto-detect format and parse file → blocks."""
    ext = filepath.suffix.lower()

    if ext == ".docx":
        return parse_docx(filepath)
    elif ext == ".html" or ext == ".htm":
        return None  # HTML files deployed as-is
    else:
        text = filepath.read_text(encoding="utf-8")
        return parse_text(text)


# ─── HTML Renderer ───

def esc(text):
    """Escape HTML special characters."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def linkify(text):
    """Convert URLs to clickable links."""
    return re.sub(r'(https?://[^\s<>\"]+)', r'<a href="\1" target="_blank">\1</a>', text)


# ─── Pretty renderer (pandoc + editorial template) ───

def has_pandoc():
    """Check if pandoc is available on PATH."""
    return shutil.which("pandoc") is not None


def render_pretty(md_source, title=None, favicon=None, source_is_path=True):
    """Render markdown → editorial HTML via pandoc + pretty template.

    md_source: file path if source_is_path=True, else raw markdown string.
    Returns the HTML string, or None if pandoc/assets are unavailable.
    """
    if not has_pandoc() or not PRETTY_CSS.exists() or not PRETTY_TEMPLATE.exists():
        return None

    title = title or "Report"
    favicon = favicon or "📄"

    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".html", delete=False, encoding="utf-8"
    ) as hdr:
        hdr.write("<style>\n")
        hdr.write(PRETTY_CSS.read_text(encoding="utf-8"))
        hdr.write("\n</style>\n")
        hdr_path = hdr.name

    try:
        cmd = [
            "pandoc",
            "-f", "gfm",
            "-t", "html",
            "--standalone",
            "--template", str(PRETTY_TEMPLATE),
            "--metadata", f"title={title}",
            "--metadata", "lang=ru",
            "--metadata", f"favicon={favicon}",
            "--include-in-header", hdr_path,
        ]
        if source_is_path:
            cmd.append(str(md_source))
            result = subprocess.run(cmd, capture_output=True, text=True)
        else:
            result = subprocess.run(
                cmd, input=md_source, capture_output=True, text=True
            )
        if result.returncode != 0:
            print(f"pandoc failed: {result.stderr}", file=sys.stderr)
            return None
        return result.stdout
    finally:
        try:
            os.unlink(hdr_path)
        except OSError:
            pass


def favicon_tag(emoji):
    """Generate inline SVG favicon from emoji."""
    if not emoji:
        return ""
    return (f'<link rel="icon" href="data:image/svg+xml,'
            f"<svg xmlns=%27http://www.w3.org/2000/svg%27 viewBox=%270 0 100 100%27>"
            f"<text y=%27.9em%27 font-size=%2790%27>{emoji}</text></svg>\">")


def render_html(blocks, title=None, favicon=None):
    """Render structured blocks → full HTML page with styles."""
    if not title:
        for btype, content in blocks:
            if btype == "h1":
                title = content[:120]
                break
        title = title or "Report"

    parts = []

    for btype, content in blocks:
        if btype in ("h1", "h2", "h3"):
            parts.append(f"<{btype}>{esc(content)}</{btype}>")

        elif btype == "p":
            text = esc(content)
            text = linkify(text)
            # Bold "Label:" at start of paragraph
            text = re.sub(
                r'^(<?)([А-Яа-яA-Za-z][А-Яа-яA-Za-z\s\-]{2,40}:)',
                r'\1<strong>\2</strong>', text
            )
            # Inline lists: "intro - item1 - item2 - item3"
            if " - " in text and len(text) > 150:
                items = text.split(" - ")
                parts.append(f"<p>{items[0]}</p>")
                parts.append("<ul>")
                for item in items[1:]:
                    item = item.strip().rstrip(".")
                    if item:
                        parts.append(f"<li>{item}</li>")
                parts.append("</ul>")
            else:
                parts.append(f"<p>{text}</p>")

        elif btype == "table":
            rows = content
            if not rows:
                continue
            max_cols = max(len(r) for r in rows)
            parts.append('<div class="table-wrap"><table>')
            for ri, row in enumerate(rows):
                while len(row) < max_cols:
                    row.append("")
                tag = "th" if ri == 0 else "td"
                parts.append("<tr>")
                for cell in row:
                    parts.append(f"<{tag}>{linkify(esc(cell))}</{tag}>")
                parts.append("</tr>")
            parts.append("</table></div>")

    body = "\n".join(parts)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    return f"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{esc(title)}</title>
{favicon_tag(favicon)}
<style>
:root {{
  --bg:#fff; --text:#1a1a2e; --muted:#6b7280;
  --accent:#2563eb; --border:#e5e7eb;
  --th-bg:#f1f5f9; --stripe:#f8fafc;
}}
@media(prefers-color-scheme:dark){{
  :root {{
    --bg:#0f172a; --text:#e2e8f0; --muted:#94a3b8;
    --accent:#60a5fa; --border:#334155;
    --th-bg:#1e293b; --stripe:#1e293b;
  }}
}}
*{{margin:0;padding:0;box-sizing:border-box}}
body{{
  font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;
  line-height:1.7;color:var(--text);background:var(--bg);
  max-width:1100px;margin:0 auto;padding:2rem 1.5rem;
}}
h1{{font-size:1.8rem;margin-bottom:1.5rem;line-height:1.3;
    border-bottom:2px solid var(--accent);padding-bottom:.5rem}}
h2{{font-size:1.3rem;margin:2rem 0 .8rem;color:var(--accent)}}
h3{{font-size:1.1rem;margin:1.5rem 0 .6rem}}
p{{margin-bottom:.8rem}}
ul{{margin:.5rem 0 1rem 1.5rem;list-style:disc}}
li{{margin-bottom:.3rem}}
a{{color:var(--accent);text-decoration:none}}
a:hover{{text-decoration:underline}}
.table-wrap{{
  overflow-x:auto;margin:1.5rem 0;
  border:1px solid var(--border);border-radius:8px
}}
table{{width:100%;border-collapse:collapse;font-size:.82rem}}
th{{
  background:var(--th-bg);font-weight:600;text-align:left;
  padding:.7rem .8rem;border-bottom:2px solid var(--border);
  position:sticky;top:0
}}
td{{padding:.6rem .8rem;border-bottom:1px solid var(--border);vertical-align:top}}
tr:nth-child(even) td{{background:var(--stripe)}}
strong{{font-weight:600}}
.footer{{
  margin-top:3rem;padding-top:1rem;
  border-top:1px solid var(--border);
  color:var(--muted);font-size:.8rem
}}
</style>
</head>
<body>
{body}
<div class="footer">Published {now}</div>
</body>
</html>"""


# ─── Slugify ───

TRANSLIT = {
    'а':'a','б':'b','в':'v','г':'g','д':'d','е':'e','ё':'yo','ж':'zh',
    'з':'z','и':'i','й':'y','к':'k','л':'l','м':'m','н':'n','о':'o',
    'п':'p','р':'r','с':'s','т':'t','у':'u','ф':'f','х':'kh','ц':'ts',
    'ч':'ch','ш':'sh','щ':'sch','ъ':'','ы':'y','ь':'','э':'e','ю':'yu','я':'ya',
}

def slugify(name):
    """Convert any string to CF-compatible project slug (a-z0-9-)."""
    slug = ""
    for ch in name.lower():
        if ch in TRANSLIT:
            slug += TRANSLIT[ch]
        elif ch.isalnum() or ch == "-":
            slug += ch
        elif ch in " _.,;:":
            slug += "-"
    slug = re.sub(r'-+', '-', slug).strip('-')
    return slug[:63] or f"report-{datetime.now().strftime('%Y%m%d')}"


# ─── Deploy ───

def ensure_project(project_name, env):
    """Create CF Pages project if it doesn't exist yet."""
    subprocess.run(
        ["wrangler", "pages", "project", "create", project_name, "--production-branch", "main"],
        capture_output=True, text=True, env=env, cwd="/tmp",
    )


def deploy(project_name, html_content):
    """Deploy HTML to Cloudflare Pages via wrangler. Returns deploy URL."""
    with tempfile.TemporaryDirectory() as tmpdir:
        (Path(tmpdir) / "index.html").write_text(html_content, encoding="utf-8")

        env = os.environ.copy()
        env["CLOUDFLARE_API_TOKEN"] = CF_API_TOKEN
        env["CLOUDFLARE_ACCOUNT_ID"] = CF_ACCOUNT_ID

        # Auto-create project on first deploy
        ensure_project(project_name, env)

        result = subprocess.run(
            ["wrangler", "pages", "deploy", tmpdir, "--project-name", project_name],
            capture_output=True, text=True, env=env, cwd="/tmp",
        )

        output = result.stdout + result.stderr

        url_match = re.search(r'https://[a-z0-9-]+\.[a-z0-9-]+\.pages\.dev', output)
        if url_match:
            return url_match.group(0)

        if result.returncode != 0:
            print(output, file=sys.stderr)
            return None

        return f"https://{project_name}.pages.dev"


# ─── CLI ───

def main():
    parser = argparse.ArgumentParser(
        description="Publish any document to Cloudflare Pages",
        epilog="Examples:\n"
               "  python3 publish.py report.docx\n"
               "  python3 publish.py notes.md --name my-notes\n"
               '  python3 publish.py page.html --name landing\n'
               '  echo "Hello" | python3 publish.py --stdin --name hello',
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("file", nargs="?", help="File to publish (.docx .txt .md .html)")
    parser.add_argument("--stdin", action="store_true", help="Read content from stdin")
    parser.add_argument("--name", help="Project name for URL (auto from filename if omitted)")
    parser.add_argument("--title", help="Page title (auto from content if omitted)")
    parser.add_argument("--favicon", help="Emoji for browser tab favicon (e.g. 🎨)")
    parser.add_argument("--legacy", action="store_true",
                        help="Force the legacy hand-rolled renderer (blue theme) "
                             "instead of the editorial pandoc theme")
    parser.add_argument("--html-only", action="store_true", help="Save HTML locally, don't deploy")
    args = parser.parse_args()

    if not args.file and not args.stdin:
        parser.print_help()
        sys.exit(1)

    use_pretty = not args.legacy and has_pandoc()

    # ── Read input ──
    if args.stdin:
        raw = sys.stdin.read()
        is_html = raw.strip().lower().startswith(("<!doctype", "<html"))
        if is_html:
            html = raw
            if args.favicon and "<head>" in html:
                html = html.replace("<head>", f"<head>\n{favicon_tag(args.favicon)}", 1)
        elif use_pretty:
            html = render_pretty(raw, title=args.title, favicon=args.favicon,
                                 source_is_path=False)
            if html is None:
                blocks = parse_text(raw)
                html = render_html(blocks, title=args.title, favicon=args.favicon)
            else:
                print("  editorial theme (pandoc)")
        else:
            blocks = parse_text(raw)
            html = render_html(blocks, title=args.title, favicon=args.favicon)
        project_name = slugify(args.name or f"page-{datetime.now().strftime('%Y%m%d-%H%M')}")
    else:
        filepath = Path(args.file)
        if not filepath.exists():
            print(f"File not found: {filepath}", file=sys.stderr)
            sys.exit(1)

        ext = filepath.suffix.lower()
        project_name = slugify(args.name or filepath.stem)

        print(f"Parsing: {filepath.name}")

        if ext in (".html", ".htm"):
            html = filepath.read_text(encoding="utf-8")
            if args.favicon and "<head>" in html:
                html = html.replace("<head>", f"<head>\n{favicon_tag(args.favicon)}", 1)
            print("  HTML — deploy as-is")
        elif ext in (".md", ".txt") and use_pretty:
            title = args.title or filepath.stem.replace("_", " ").replace("-", " ").title()
            html = render_pretty(filepath, title=title, favicon=args.favicon)
            if html is None:
                text = filepath.read_text(encoding="utf-8")
                blocks = parse_text(text)
                print(f"  {len(blocks)} blocks (legacy fallback)")
                html = render_html(blocks, title=args.title, favicon=args.favicon)
            else:
                print("  editorial theme (pandoc)")
        else:
            if ext == ".docx":
                blocks = parse_docx(filepath)
            else:
                text = filepath.read_text(encoding="utf-8")
                blocks = parse_text(text)

            print(f"  {len(blocks)} blocks")
            html = render_html(blocks, title=args.title, favicon=args.favicon)

    # ── Output ──
    if args.html_only:
        if args.file:
            out = Path(args.file).with_suffix(".html")
        else:
            out = Path(f"{project_name}.html")
        out.write_text(html, encoding="utf-8")
        print(f"HTML saved: {out}")
        return

    if not CF_ACCOUNT_ID or not CF_API_TOKEN:
        print("Error: CF_ACCOUNT_ID and CF_API_TOKEN not found in .env", file=sys.stderr)
        sys.exit(1)

    print(f"Deploying → {project_name}")
    url = deploy(project_name, html)

    if url:
        pages_url = f"https://{project_name}.pages.dev"
        print(f"\n  Deploy:    {url}")
        print(f"  Permanent: {pages_url}")
    else:
        print("Deploy failed!", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
